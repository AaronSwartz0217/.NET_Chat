#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
根据项目内容更新实验报告文档
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def update_report():
    # 创建新文档
    doc = Document()
    
    # 设置字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    
    # 项目信息
    project_info = {
        "project_name": "社区论坛与即时通讯系统",
        "modules": [
            {"name": "用户认证模块", "desc": "实现用户注册、登录、密码修改、JWT令牌管理"},
            {"name": "论坛系统模块", "desc": "支持帖子发布、评论、点赞、收藏、浏览统计"},
            {"name": "即时聊天模块", "desc": "实现私聊、群聊、消息撤回、已读标记"},
            {"name": "资讯中心模块", "desc": "内嵌腾讯新闻，提供每日资讯浏览"},
            {"name": "学生档案模块", "desc": "用户资料完善、学生信息管理"}
        ],
        "tech_stack": [
            ("前端框架", "Avalonia UI 11.x", "跨平台桌面应用框架"),
            ("后端框架", "ASP.NET Core 9.0", "高性能Web API框架"),
            ("数据库", "MySQL 8.0+", "关系型数据库"),
            ("ORM框架", "SqlSugar 5.x", "数据库操作框架"),
            ("认证方式", "JWT", "JSON Web Token认证"),
            ("实时通信", "WebSocket", "即时聊天功能")
        ],
        "api_list": [
            ("POST", "/api/auth/login", "用户登录"),
            ("POST", "/api/auth/register", "用户注册"),
            ("POST", "/api/auth/refresh", "刷新令牌"),
            ("GET", "/api/posts", "获取帖子列表"),
            ("POST", "/api/posts", "创建帖子"),
            ("GET", "/api/posts/{id}", "获取帖子详情"),
            ("GET", "/api/channels", "获取会话列表"),
            ("POST", "/api/channels/{id}/messages", "发送消息")
        ],
        "test_cases": [
            ("用户注册", "成功创建新用户"),
            ("用户登录", "成功获取JWT令牌"),
            ("发布帖子", "帖子保存到数据库"),
            ("浏览帖子", "分页显示帖子列表"),
            ("点赞帖子", "点赞数正确增加"),
            ("发送消息", "消息实时推送"),
            ("资讯浏览", "内嵌腾讯新闻正常显示")
        ]
    }
    
    print("🔄 创建封面...")
    # 封面
    title1 = doc.add_heading("2025-2026-2", level=1)
    title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title2 = doc.add_heading("信息系统设计与实践 & 软件工程实践", level=1)
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title3 = doc.add_heading("报告书", level=1)
    title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for _ in range(8):
        doc.add_paragraph("")
    
    # 班级信息表格
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    table.cell(0, 0).text = "班级"
    table.cell(0, 1).text = "学号"
    table.cell(0, 2).text = "姓名"
    table.cell(1, 0).text = "计233"
    table.cell(1, 1).text = "1111111111"
    table.cell(1, 2).text = "成员1"
    table.cell(2, 0).text = "计233"
    table.cell(2, 1).text = ""
    table.cell(2, 2).text = "成员2"
    table.cell(3, 0).text = "计233"
    table.cell(3, 1).text = ""
    table.cell(3, 2).text = "成员3"
    table.cell(4, 0).text = ""
    table.cell(4, 1).text = ""
    table.cell(4, 2).text = "成员4"
    table.cell(5, 0).text = "任课教师"
    table.cell(5, 1).text = "陈晓勇"
    table.cell(5, 2).text = ""
    
    # 报告书写要求
    print("🔄 添加报告书写要求...")
    doc.add_page_break()
    doc.add_heading("报告书写要求", level=1)
    doc.add_paragraph("1. 报告书需按格式要求完整填写，不得留空或改动版式：")
    doc.add_paragraph("   一）实验任务描述本软件实现的内容和效果；")
    doc.add_paragraph("   二）技术路线请写出所用到的技术简介（包括用到的组件包等）；")
    doc.add_paragraph("   三）系统实现给出所有设计的主操作界面、操作简介等，包括关键性代码（代码太多则无需写全部，挑关键代码即可），代码粘贴时注意排版和格式（可使用某些网站提供的工具进行格式化处理），适当控制字体和行间距，保持页面美观；")
    doc.add_paragraph("2. 请勿打乱报告书整体板式，各部分内容可根据实际长度自行调整内容高度；")
    doc.add_paragraph("3. 代码必须配有适当的注释，代码书写规范；")
    doc.add_paragraph("4. 严禁直接抄袭他人成果和报告内容；")
    doc.add_paragraph("5. 本报告书最后提交Word电子版。")
    
    # 实验任务
    print("🔄 添加实验任务章节...")
    doc.add_page_break()
    doc.add_heading("一、实验任务", level=1)
    doc.add_paragraph("1.1 项目概述")
    doc.add_paragraph(f"    本项目开发了一个{project_info['project_name']}，旨在为用户提供一个集社交、资讯、交流于一体的综合性平台。")
    doc.add_paragraph("1.2 实现内容")
    for module in project_info["modules"]:
        doc.add_paragraph(f"    • {module['name']}: {module['desc']}")
    doc.add_paragraph("1.3 预期效果")
    doc.add_paragraph("    ✅ 实现完整的用户认证体系")
    doc.add_paragraph("    ✅ 支持论坛帖子的CRUD操作")
    doc.add_paragraph("    ✅ 提供实时聊天功能")
    doc.add_paragraph("    ✅ 集成资讯浏览模块")
    doc.add_paragraph("    ✅ 支持Linux服务器部署")
    doc.add_paragraph("    ✅ 局域网内多设备访问")
    
    # 技术路线
    print("🔄 添加技术路线章节...")
    doc.add_page_break()
    doc.add_heading("二、技术路线", level=1)
    doc.add_paragraph("2.1 技术架构")
    doc.add_paragraph("")
    doc.add_paragraph("    ┌─────────────────────────────────────────┐")
    doc.add_paragraph("    │           客户端 (Avalonia UI)          │")
    doc.add_paragraph("    ├─────────────────────────────────────────┤")
    doc.add_paragraph("    │           后端服务 (ASP.NET Core)       │")
    doc.add_paragraph("    ├─────────────────────────────────────────┤")
    doc.add_paragraph("    │           数据库 (MySQL)               │")
    doc.add_paragraph("    └─────────────────────────────────────────┘")
    doc.add_paragraph("")
    doc.add_paragraph("2.2 核心技术栈")
    doc.add_paragraph("")
    for tech in project_info["tech_stack"]:
        doc.add_paragraph(f"    • {tech[0]}: {tech[1]} - {tech[2]}")
    
    # 系统实现
    print("🔄 添加系统实现章节...")
    doc.add_page_break()
    doc.add_heading("三、系统实现", level=1)
    doc.add_paragraph("3.1 主操作界面")
    doc.add_paragraph("(1) 登录界面")
    doc.add_paragraph("    ┌─────────────────────────────┐")
    doc.add_paragraph("    │         系统登录            │")
    doc.add_paragraph("    ├─────────────────────────────┤")
    doc.add_paragraph("    │  用户名: [___________]      │")
    doc.add_paragraph("    │  密码:   [___________]      │")
    doc.add_paragraph("    │                             │")
    doc.add_paragraph("    │     [ 登录 ]  [ 注册 ]      │")
    doc.add_paragraph("    └─────────────────────────────┘")
    doc.add_paragraph("")
    doc.add_paragraph("3.2 数据库设计")
    doc.add_paragraph("    系统采用MySQL数据库，包含以下核心表：")
    doc.add_paragraph("    • users - 用户表（用户ID、用户名、密码、邮箱等）")
    doc.add_paragraph("    • posts - 帖子表（帖子ID、作者ID、标题、内容、创建时间等）")
    doc.add_paragraph("    • comments - 评论表（评论ID、帖子ID、用户ID、内容等）")
    doc.add_paragraph("    • messages - 消息表（消息ID、会话ID、发送者ID、内容等）")
    doc.add_paragraph("")
    doc.add_paragraph("3.3 API接口设计")
    doc.add_paragraph("")
    
    table = doc.add_table(rows=len(project_info["api_list"])+1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "方法"
    hdr[1].text = "接口"
    hdr[2].text = "说明"
    for i, api in enumerate(project_info["api_list"]):
        row = table.rows[i+1].cells
        row[0].text = api[0]
        row[1].text = api[1]
        row[2].text = api[2]
    
    # 系统部署
    print("🔄 添加系统部署章节...")
    doc.add_page_break()
    doc.add_heading("四、系统部署", level=1)
    doc.add_paragraph("4.1 本地开发环境")
    doc.add_paragraph("    # 启动后端服务")
    doc.add_paragraph("    cd Chat.Server")
    doc.add_paragraph("    dotnet run --urls \"http://0.0.0.0:5002\"")
    doc.add_paragraph("")
    doc.add_paragraph("4.2 Linux服务器部署")
    doc.add_paragraph("    1. 上传部署包到服务器")
    doc.add_paragraph("    2. 运行安装脚本：sudo ./install-debian.sh")
    doc.add_paragraph("    3. 启动服务：sudo systemctl start chat-server")
    
    # 功能测试
    print("🔄 添加功能测试章节...")
    doc.add_page_break()
    doc.add_heading("五、功能测试", level=1)
    table = doc.add_table(rows=len(project_info["test_cases"])+1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "测试项"
    hdr[1].text = "预期结果"
    for i, test in enumerate(project_info["test_cases"]):
        row = table.rows[i+1].cells
        row[0].text = test[0]
        row[1].text = test[1]
    
    # 总结
    print("🔄 添加总结章节...")
    doc.add_page_break()
    doc.add_heading("六、总结", level=1)
    doc.add_paragraph(f"    本项目成功实现了{project_info['project_name']}，涵盖用户认证、论坛管理、即时聊天、资讯浏览等核心功能。")
    doc.add_paragraph("    系统采用前后端分离架构，支持跨平台部署，具备良好的扩展性和安全性。")
    
    # 保存文档
    output_path = "计科实践报告_完整版.docx"
    doc.save(output_path)
    print(f"✅ 文档已创建完成: {output_path}")
    print(f"📄 总段落数: {len(doc.paragraphs)}")
    print(f"📊 表格数: {len(doc.tables)}")

if __name__ == "__main__":
    update_report()