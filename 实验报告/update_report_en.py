#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def update_report():
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style.font.size = Pt(12)
    
    project_info = {
        "project_name": "Community Forum and Instant Messaging System",
        "modules": [
            {"name": "User Authentication", "desc": "User registration, login, password change, JWT token management"},
            {"name": "Forum System", "desc": "Post creation, comments, likes, favorites, view statistics"},
            {"name": "Instant Messaging", "desc": "Private chat, group chat, message recall, read marks"},
            {"name": "News Center", "desc": "Embedded Tencent News for daily news browsing"},
            {"name": "Student Profile", "desc": "User profile completion, student information management"}
        ],
        "tech_stack": [
            ("Frontend Framework", "Avalonia UI 11.x", "Cross-platform desktop framework"),
            ("Backend Framework", "ASP.NET Core 9.0", "High-performance Web API"),
            ("Database", "MySQL 8.0+", "Relational database"),
            ("ORM", "SqlSugar 5.x", "Database ORM framework"),
            ("Authentication", "JWT", "JSON Web Token"),
            ("Realtime", "WebSocket", "Instant messaging")
        ],
        "api_list": [
            ("POST", "/api/auth/login", "User login"),
            ("POST", "/api/auth/register", "User registration"),
            ("POST", "/api/auth/refresh", "Refresh token"),
            ("GET", "/api/posts", "Get post list"),
            ("POST", "/api/posts", "Create post"),
            ("GET", "/api/posts/{id}", "Get post detail"),
            ("GET", "/api/channels", "Get channel list"),
            ("POST", "/api/channels/{id}/messages", "Send message")
        ],
        "test_cases": [
            ("User Registration", "Successfully create new user"),
            ("User Login", "Successfully get JWT token"),
            ("Create Post", "Post saved to database"),
            ("View Posts", "Paginated post list displayed"),
            ("Like Post", "Like count increased correctly"),
            ("Send Message", "Message pushed in realtime"),
            ("News Browse", "Embedded news displayed correctly")
        ]
    }
    
    print("Creating cover page...")
    title1 = doc.add_heading("2025-2026-2", level=1)
    title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    title2 = doc.add_heading("Information System Design & Software Engineering Practice", level=1)
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    title3 = doc.add_heading("Report", level=1)
    title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for _ in range(8):
        doc.add_paragraph("")
    
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    table.cell(0, 0).text = "Class"
    table.cell(0, 1).text = "Student ID"
    table.cell(0, 2).text = "Name"
    table.cell(1, 0).text = "CS233"
    table.cell(1, 1).text = "1111111111"
    table.cell(1, 2).text = "Member 1"
    table.cell(2, 0).text = "CS233"
    table.cell(2, 1).text = ""
    table.cell(2, 2).text = "Member 2"
    table.cell(3, 0).text = "CS233"
    table.cell(3, 1).text = ""
    table.cell(3, 2).text = "Member 3"
    table.cell(4, 0).text = ""
    table.cell(4, 1).text = ""
    table.cell(4, 2).text = "Member 4"
    table.cell(5, 0).text = "Teacher"
    table.cell(5, 1).text = "Chen Xiaoyong"
    table.cell(5, 2).text = ""
    
    print("Adding requirements section...")
    doc.add_page_break()
    doc.add_heading("Report Requirements", level=1)
    doc.add_paragraph("1. The report must be completed according to the format requirements.")
    doc.add_paragraph("2. Do not disrupt the overall layout of the report.")
    doc.add_paragraph("3. Code must be properly commented and follow coding standards.")
    doc.add_paragraph("4. Plagiarism is strictly prohibited.")
    doc.add_paragraph("5. Submit the report in Word format.")
    
    print("Adding task section...")
    doc.add_page_break()
    doc.add_heading("1. Project Task", level=1)
    doc.add_paragraph("1.1 Project Overview")
    doc.add_paragraph("    This project develops a " + project_info["project_name"] + ", providing a comprehensive platform for social interaction, news, and communication.")
    doc.add_paragraph("1.2 Implementation Content")
    for module in project_info["modules"]:
        doc.add_paragraph("    • " + module["name"] + ": " + module["desc"])
    doc.add_paragraph("1.3 Expected Results")
    doc.add_paragraph("    • Complete user authentication system")
    doc.add_paragraph("    • Forum CRUD operations")
    doc.add_paragraph("    • Real-time chat functionality")
    doc.add_paragraph("    • News browsing module")
    doc.add_paragraph("    • Linux server deployment support")
    
    print("Adding technology section...")
    doc.add_page_break()
    doc.add_heading("2. Technical Route", level=1)
    doc.add_paragraph("2.1 System Architecture")
    doc.add_paragraph("")
    doc.add_paragraph("    ┌─────────────────────────────────────────┐")
    doc.add_paragraph("    │           Client (Avalonia UI)          │")
    doc.add_paragraph("    ├─────────────────────────────────────────┤")
    doc.add_paragraph("    │           Backend (ASP.NET Core)        │")
    doc.add_paragraph("    ├─────────────────────────────────────────┤")
    doc.add_paragraph("    │           Database (MySQL)              │")
    doc.add_paragraph("    └─────────────────────────────────────────┘")
    doc.add_paragraph("")
    doc.add_paragraph("2.2 Core Technologies")
    doc.add_paragraph("")
    for tech in project_info["tech_stack"]:
        doc.add_paragraph("    • " + tech[0] + ": " + tech[1] + " - " + tech[2])
    
    print("Adding implementation section...")
    doc.add_page_break()
    doc.add_heading("3. System Implementation", level=1)
    doc.add_paragraph("3.1 Main Interface")
    doc.add_paragraph("(1) Login Interface")
    doc.add_paragraph("    ┌─────────────────────────────┐")
    doc.add_paragraph("    │         System Login        │")
    doc.add_paragraph("    ├─────────────────────────────┤")
    doc.add_paragraph("    │  Username: [___________]    │")
    doc.add_paragraph("    │  Password: [___________]    │")
    doc.add_paragraph("    │                             │")
    doc.add_paragraph("    │    [ Login ]  [ Register ]  │")
    doc.add_paragraph("    └─────────────────────────────┘")
    doc.add_paragraph("")
    doc.add_paragraph("3.2 Database Design")
    doc.add_paragraph("    The system uses MySQL database with the following core tables:")
    doc.add_paragraph("    • users - User information")
    doc.add_paragraph("    • posts - Forum posts")
    doc.add_paragraph("    • comments - Comments")
    doc.add_paragraph("    • messages - Chat messages")
    doc.add_paragraph("")
    doc.add_paragraph("3.3 API Interface Design")
    doc.add_paragraph("")
    
    table = doc.add_table(rows=len(project_info["api_list"])+1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "Method"
    hdr[1].text = "Endpoint"
    hdr[2].text = "Description"
    for i, api in enumerate(project_info["api_list"]):
        row = table.rows[i+1].cells
        row[0].text = api[0]
        row[1].text = api[1]
        row[2].text = api[2]
    
    print("Adding deployment section...")
    doc.add_page_break()
    doc.add_heading("4. System Deployment", level=1)
    doc.add_paragraph("4.1 Local Development")
    doc.add_paragraph("    # Start backend service")
    doc.add_paragraph("    cd Chat.Server")
    doc.add_paragraph("    dotnet run --urls \"http://0.0.0.0:5002\"")
    doc.add_paragraph("")
    doc.add_paragraph("4.2 Linux Server Deployment")
    doc.add_paragraph("    1. Upload deployment package to server")
    doc.add_paragraph("    2. Run installation script: sudo ./install-debian.sh")
    doc.add_paragraph("    3. Start service: sudo systemctl start chat-server")
    
    print("Adding testing section...")
    doc.add_page_break()
    doc.add_heading("5. Functional Testing", level=1)
    table = doc.add_table(rows=len(project_info["test_cases"])+1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "Test Item"
    hdr[1].text = "Expected Result"
    for i, test in enumerate(project_info["test_cases"]):
        row = table.rows[i+1].cells
        row[0].text = test[0]
        row[1].text = test[1]
    
    print("Adding summary section...")
    doc.add_page_break()
    doc.add_heading("6. Summary", level=1)
    doc.add_paragraph("    This project successfully implemented " + project_info["project_name"] + ", covering user authentication, forum management, instant messaging, and news browsing.")
    doc.add_paragraph("    The system uses a decoupled architecture, supports cross-platform deployment, and has good scalability and security.")
    
    output_path = "Experiment_Report_Full.docx"
    doc.save(output_path)
    print("Report saved successfully: " + output_path)

if __name__ == "__main__":
    update_report()