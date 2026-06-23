#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
实验报告处理脚本
读取现有Word文档并根据项目内容进行修改
"""

import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("请先安装依赖:")
    print("pip install python-docx")
    sys.exit(1)

# 项目信息常量
PROJECT_INFO = {
    "project_name": "社区论坛与即时通讯系统",
    "version": "1.0.0",
    "team": "计233班",
    "teacher": "陈晓勇",
    "date": "2025-2026-2",
    
    "modules": [
        {"name": "用户认证", "description": "用户注册、登录、密码修改、JWT令牌管理"},
        {"name": "论坛系统", "description": "帖子发布、评论、点赞、收藏、浏览统计"},
        {"name": "即时聊天", "description": "私聊、群聊、消息撤回、已读标记"},
        {"name": "资讯中心", "description": "内嵌腾讯新闻，提供每日资讯浏览"},
        {"name": "学生档案", "description": "用户资料完善、学生信息管理"}
    ],
    
    "tech_stack": [
        {"category": "前端框架", "tech": "Avalonia UI", "version": "11.x", "desc": "跨平台桌面应用框架"},
        {"category": "后端框架", "tech": "ASP.NET Core", "version": "9.0", "desc": "高性能Web API框架"},
        {"category": "数据库", "tech": "MySQL", "version": "8.0+", "desc": "关系型数据库"},
        {"category": "ORM", "tech": "SqlSugar", "version": "5.x", "desc": "数据库操作框架"},
        {"category": "认证", "tech": "JWT", "version": "-", "desc": "JSON Web Token认证"},
        {"category": "实时通信", "tech": "WebSocket", "version": "-", "desc": "即时聊天功能"}
    ],
    
    "api_list": [
        {"module": "认证", "path": "/api/auth/login", "method": "POST", "desc": "用户登录"},
        {"module": "认证", "path": "/api/auth/register", "method": "POST", "desc": "用户注册"},
        {"module": "认证", "path": "/api/auth/refresh", "method": "POST", "desc": "刷新令牌"},
        {"module": "帖子", "path": "/api/posts", "method": "GET", "desc": "获取帖子列表"},
        {"module": "帖子", "path": "/api/posts", "method": "POST", "desc": "创建帖子"},
        {"module": "帖子", "path": "/api/posts/{id}", "method": "GET", "desc": "获取帖子详情"},
        {"module": "聊天", "path": "/api/channels", "method": "GET", "desc": "获取会话列表"},
        {"module": "聊天", "path": "/api/channels/{id}/messages", "method": "GET", "desc": "获取消息历史"}
    ],
    
    "test_cases": [
        {"item": "用户注册", "expected": "成功创建新用户，返回用户信息"},
        {"item": "用户登录", "expected": "成功获取JWT令牌"},
        {"item": "发布帖子", "expected": "帖子保存到数据库"},
        {"item": "浏览帖子", "expected": "分页显示帖子列表"},
        {"item": "点赞帖子", "expected": "点赞数增加"},
        {"item": "发送消息", "expected": "消息实时推送"},
        {"item": "资讯浏览", "expected": "内嵌腾讯新闻正常显示"}
    ]
}

def read_document(file_path):
    """读取Word文档内容"""
    if not os.path.exists(file_path):
        print(f"❌ 文件不存在: {file_path}")
        return None
    
    try:
        doc = Document(file_path)
        print(f"✅ 成功读取文档: {file_path}")
        print(f"📄 段落数: {len(doc.paragraphs)}")
        print(f"📊 表格数: {len(doc.tables)}")
        return doc
    except Exception as e:
        print(f"❌ 读取文档失败: {e}")
        return None

def display_content(doc):
    """显示文档内容"""
    print("\n" + "="*60)
    print("文档内容预览")
    print("="*60)
    
    for i, paragraph in enumerate(doc.paragraphs[:20]):  # 只显示前20段
        text = paragraph.text.strip()
        if text:
            print(f"{i+1:3d}. {text[:100]}{'...' if len(text) > 100 else ''}")
    
    if len(doc.paragraphs) > 20:
        print(f"... 还有 {len(doc.paragraphs) - 20} 段未显示")

def find_and_replace(doc, old_text, new_text):
    """查找并替换文本"""
    count = 0
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text)
            count += 1
    print(f"🔄 替换完成，共替换 {count} 处")
    return count

def update_section(doc, section_title, new_content):
    """更新指定章节的内容"""
    # 查找章节位置
    start_idx = None
    end_idx = None
    
    for i, paragraph in enumerate(doc.paragraphs):
        if section_title in paragraph.text:
            start_idx = i
            # 查找下一个标题或文档末尾
            for j in range(i+1, len(doc.paragraphs)):
                next_paragraph = doc.paragraphs[j]
                # 判断是否是下一个标题（通常是加粗或标题样式）
                if (next_paragraph.style.name.startswith('Heading') or 
                    (next_paragraph.runs and next_paragraph.runs[0].bold)):
                    end_idx = j
                    break
            if end_idx is None:
                end_idx = len(doc.paragraphs)
            break
    
    if start_idx is None:
        print(f"❌ 未找到章节: {section_title}")
        return False
    
    # 删除旧内容
    for i in range(end_idx - 1, start_idx, -1):
        p = doc.paragraphs[i]
        doc._body.remove(p._element)
    
    # 添加新内容
    current_idx = start_idx + 1
    for content in new_content:
        paragraph = doc.add_paragraph(content)
        # 插入到正确位置
        doc.paragraphs.insert(current_idx, paragraph)
        current_idx += 1
    
    print(f"✅ 成功更新章节: {section_title}")
    return True

def generate_module_content():
    """生成模块内容"""
    content = []
    content.append("1.2 实现内容")
    content.append("")
    for module in PROJECT_INFO["modules"]:
        content.append(f"   • {module['name']}: {module['description']}")
    content.append("")
    content.append("1.3 预期效果")
    content.append("")
    content.append("    ✅ 实现完整的用户认证体系")
    content.append("    ✅ 支持论坛帖子的CRUD操作")
    content.append("    ✅ 提供实时聊天功能")
    content.append("    ✅ 集成资讯浏览模块")
    content.append("    ✅ 支持Linux服务器部署")
    content.append("    ✅ 局域网内多设备访问")
    return content

def generate_tech_content():
    """生成技术路线内容"""
    content = []
    content.append("2.1 技术架构")
    content.append("")
    content.append("    ┌─────────────────────────────────────────┐")
    content.append("    │           客户端 (Avalonia UI)          │")
    content.append("    ├─────────────────────────────────────────┤")
    content.append("    │           后端服务 (ASP.NET Core)       │")
    content.append("    ├─────────────────────────────────────────┤")
    content.append("    │           数据库 (MySQL)               │")
    content.append("    └─────────────────────────────────────────┘")
    content.append("")
    content.append("2.2 核心技术栈")
    content.append("")
    for tech in PROJECT_INFO["tech_stack"]:
        content.append(f"   • {tech['category']}: {tech['tech']} v{tech['version']} - {tech['desc']}")
    return content

def generate_api_content():
    """生成API接口内容"""
    content = []
    content.append("3.3 API接口设计")
    content.append("")
    content.append("    | 模块 | 接口 | 方法 | 说明 |")
    content.append("    |------|------|------|------|")
    for api in PROJECT_INFO["api_list"]:
        content.append(f"    | {api['module']} | {api['path']} | {api['method']} | {api['desc']} |")
    return content

def generate_test_content():
    """生成测试内容"""
    content = []
    content.append("五、功能测试")
    content.append("")
    content.append("    | 测试项 | 预期结果 |")
    content.append("    |--------|----------|")
    for test in PROJECT_INFO["test_cases"]:
        content.append(f"    | {test['item']} | {test['expected']} |")
    return content

def create_new_report():
    """创建新的实验报告"""
    doc = Document()
    
    # 设置字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    
    # 封面
    title1 = doc.add_heading(PROJECT_INFO["date"], level=1)
    title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    title2 = doc.add_heading("信息系统设计与实践 & 软件工程实践", level=1)
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    title3 = doc.add_heading("报告书", level=1)
    title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 空行
    for _ in range(8):
        doc.add_paragraph("")
    
    # 班级信息表格
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    table.cell(0, 0).text = "班级"
    table.cell(0, 1).text = "学号"
    table.cell(0, 2).text = "姓名"
    table.cell(1, 0).text = "计233"
    table.cell(1, 1).text = "1111111111"
    table.cell(1, 2).text = "成员1"
    table.cell(2, 0).text = "计233"
    table.cell(3, 0).text = "计233"
    table.cell(4, 2).text = "成员4"
    table.cell(5, 0).text = "任课教师"
    table.cell(5, 1).text = PROJECT_INFO["teacher"]
    
    # 报告要求
    doc.add_page_break()
    doc.add_heading("报告书写要求", level=1)
    doc.add_paragraph("1. 报告书需按格式要求完整填写，不得留空或改动版式：")
    doc.add_paragraph("   一）实验任务描述本软件实现的内容和效果；")
    doc.add_paragraph("   二）技术路线请写出所用到的技术简介（包括用到的组件包等）；")
    doc.add_paragraph("   三）系统实现给出所有设计的主操作界面、操作简介等，包括关键性代码。")
    doc.add_paragraph("2. 请勿打乱报告书整体板式，各部分内容可根据实际长度自行调整内容高度；")
    doc.add_paragraph("3. 代码必须配有适当的注释，代码书写规范；")
    doc.add_paragraph("4. 严禁直接抄袭他人成果和报告内容；")
    doc.add_paragraph("5. 本报告书最后提交Word电子版。")
    
    # 实验任务
    doc.add_page_break()
    doc.add_heading("一、实验任务", level=1)
    doc.add_paragraph("1.1 项目概述")
    doc.add_paragraph(f"    本项目开发了一个{PROJECT_INFO['project_name']}，旨在为用户提供一个集社交、资讯、交流于一体的综合性平台。")
    
    for content in generate_module_content():
        doc.add_paragraph(content)
    
    # 技术路线
    doc.add_page_break()
    doc.add_heading("二、技术路线", level=1)
    for content in generate_tech_content():
        doc.add_paragraph(content)
    
    # 系统实现
    doc.add_page_break()
    doc.add_heading("三、系统实现", level=1)
    doc.add_paragraph("3.1 主操作界面")
    doc.add_paragraph("(1) 登录界面")
    doc.add_paragraph("    ┌─────────────────────────────┐")
    doc.add_paragraph("    │         系统登录            │")
    doc.add_paragraph("    ├─────────────────────────────┤")
    doc.add_paragraph("    │  用户名: [___________]      │")
    doc.add_paragraph("    │  密码:   [___________]      │")
    doc.add_paragraph("    │                             │")
    doc.add_paragraph("    │     [ 登录 ]  [ 注册 ]      │")
    doc.add_paragraph("    └─────────────────────────────┘")
    doc.add_paragraph("")
    doc.add_paragraph("3.2 数据库设计")
    doc.add_paragraph("    系统采用MySQL数据库，包含以下核心表：")
    doc.add_paragraph("    • users - 用户表")
    doc.add_paragraph("    • posts - 帖子表")
    doc.add_paragraph("    • messages - 消息表")
    doc.add_paragraph("    • students - 学生档案表")
    
    for content in generate_api_content():
        doc.add_paragraph(content)
    
    # 系统部署
    doc.add_page_break()
    doc.add_heading("四、系统部署", level=1)
    doc.add_paragraph("4.1 本地开发环境")
    doc.add_paragraph("    运行后端服务：dotnet run --urls \"http://0.0.0.0:5002\"")
    doc.add_paragraph("    运行客户端：dotnet run (Chat.Desktop)")
    doc.add_paragraph("")
    doc.add_paragraph("4.2 Linux服务器部署")
    doc.add_paragraph("    1. 上传部署包到服务器")
    doc.add_paragraph("    2. 运行安装脚本：sudo ./install-debian.sh")
    doc.add_paragraph("    3. 启动服务：sudo systemctl start chat-server")
    
    # 功能测试
    doc.add_page_break()
    for content in generate_test_content():
        doc.add_paragraph(content)
    
    # 总结
    doc.add_page_break()
    doc.add_heading("六、总结", level=1)
    doc.add_paragraph(f"    本项目成功实现了{PROJECT_INFO['project_name']}，涵盖用户认证、论坛管理、即时聊天、资讯浏览等核心功能。")
    doc.add_paragraph("    系统采用前后端分离架构，支持跨平台部署，具备良好的扩展性和安全性。")
    
    # 保存
    output_path = "实验报告_完整版.docx"
    doc.save(output_path)
    print(f"✅ 新报告已创建: {output_path}")

def main():
    """主函数"""
    print("=" * 60)
    print("      实验报告处理脚本")
    print("=" * 60)
    print("1. 读取现有文档")
    print("2. 修改现有文档")
    print("3. 创建新报告")
    print("4. 退出")
    print("=" * 60)
    
    choice = input("请输入选择 (1/2/3/4): ")
    
    if choice == "1":
        file_path = input("请输入文档路径: ").strip()
        doc = read_document(file_path)
        if doc:
            display_content(doc)
    
    elif choice == "2":
        file_path = input("请输入文档路径: ").strip()
        doc = read_document(file_path)
        if doc:
            print("\n可选操作:")
            print("1. 更新实验任务章节")
            print("2. 更新技术路线章节")
            print("3. 更新API接口章节")
            print("4. 更新测试章节")
            sub_choice = input("请输入选择 (1/2/3/4): ")
            
            if sub_choice == "1":
                update_section(doc, "一、实验任务", generate_module_content())
            elif sub_choice == "2":
                update_section(doc, "二、技术路线", generate_tech_content())
            elif sub_choice == "3":
                update_section(doc, "三、系统实现", generate_api_content())
            elif sub_choice == "4":
                update_section(doc, "五、功能测试", generate_test_content())
            
            save_path = input("请输入保存路径 (回车覆盖原文件): ").strip() or file_path
            doc.save(save_path)
            print(f"✅ 文档已保存: {save_path}")
    
    elif choice == "3":
        create_new_report()
    
    elif choice == "4":
        print("👋 退出程序")
    
    else:
        print("❌ 无效选择")

if __name__ == "__main__":
    main()