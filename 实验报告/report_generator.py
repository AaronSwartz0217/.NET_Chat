#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
实验报告自动生成脚本
根据项目内容自动更新Word文档
"""

import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION
except ImportError:
    print("请先安装python-docx库: pip install python-docx")
    sys.exit(1)

class ReportGenerator:
    """实验报告生成器"""
    
    def __init__(self, template_path=None):
        self.template_path = template_path
        if template_path and os.path.exists(template_path):
            self.doc = Document(template_path)
        else:
            self.doc = Document()
        
        # 设置默认字体
        self.style = self.doc.styles['Normal']
        self.style.font.name = '宋体'
        self.style.font.size = Pt(12)
    
    def add_title(self, text, level=1):
        """添加标题"""
        if level == 1:
            heading = self.doc.add_heading(text, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            self.doc.add_heading(text, level=level)
    
    def add_paragraph(self, text, style=None):
        """添加段落"""
        p = self.doc.add_paragraph(text, style=style)
        p.paragraph_format.line_spacing = 1.5
        return p
    
    def add_table(self, data, headers=None):
        """添加表格"""
        rows = len(data)
        cols = len(data[0]) if data else 0
        
        table = self.doc.add_table(rows=rows + (1 if headers else 0), cols=cols)
        table.style = 'Table Grid'
        
        # 添加表头
        if headers:
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
        
        # 添加数据
        start_row = 1 if headers else 0
        for i, row in enumerate(data):
            row_cells = table.rows[i + start_row].cells
            for j, cell in enumerate(row):
                row_cells[j].text = str(cell)
        
        return table
    
    def add_section(self):
        """添加新章节"""
        self.doc.add_section(WD_SECTION.NEW_PAGE)
    
    def save(self, output_path):
        """保存文档"""
        self.doc.save(output_path)
        print(f"文档已保存: {output_path}")

def generate_report():
    """生成实验报告"""
    # 创建报告生成器
    report = ReportGenerator()
    
    # ===== 封面 =====
    report.add_title("2025-2026-2")
    report.add_title("信息系统设计与实践 & 软件工程实践")
    report.add_title("报告书")
    
    # 空行
    for _ in range(8):
        report.add_paragraph("")
    
    # 班级信息表格
    info_table = [
        ["班级", "学号", "姓名"],
        ["计233", "1111111111", "成员1"],
        ["计233", "", "成员2"],
        ["计233", "", "成员3"],
        ["", "", "成员4"],
        ["任课教师", "陈晓勇", ""]
    ]
    report.add_table(info_table)
    
    # ===== 报告书写要求 =====
    report.add_section()
    report.add_title("报告书写要求", level=1)
    
    report.add_paragraph("1. 报告书需按格式要求完整填写，不得留空或改动版式：")
    report.add_paragraph("   一）实验任务描述本软件实现的内容和效果；")
    report.add_paragraph("   二）技术路线请写出所用到的技术简介（包括用到的组件包等）；")
    report.add_paragraph("   三）系统实现给出所有设计的主操作界面、操作简介等，包括关键性代码（代码太多则无需写全部，挑关键代码即可），代码粘贴时注意排版和格式（可使用某些网站提供的工具进行格式化处理），适当控制字体和行间距，保持页面美观；")
    report.add_paragraph("2. 请勿打乱报告书整体板式，各部分内容可根据实际长度自行调整内容高度；")
    report.add_paragraph("3. 代码必须配有适当的注释，代码书写规范；")
    report.add_paragraph("4. 严禁直接抄袭他人成果和报告内容；")
    report.add_paragraph("5. 本报告书最后提交Word电子版。")
    
    # ===== 一、实验任务 =====
    report.add_section()
    report.add_title("一、实验任务", level=1)
    
    report.add_paragraph("1.1 项目概述")
    report.add_paragraph("    本项目开发了一个综合性社区论坛与即时通讯系统，旨在为用户提供一个集社交、资讯、交流于一体的综合性平台。")
    
    report.add_paragraph("1.2 实现内容")
    content_table = [
        ["模块", "功能描述"],
        ["用户认证", "用户注册、登录、密码修改、JWT令牌管理"],
        ["论坛系统", "帖子发布、评论、点赞、收藏、浏览统计"],
        ["即时聊天", "私聊、群聊、消息撤回、已读标记"],
        ["资讯中心", "内嵌腾讯新闻，提供每日资讯浏览"],
        ["学生档案", "用户资料完善、学生信息管理"]
    ]
    report.add_table(content_table)
    
    report.add_paragraph("1.3 预期效果")
    report.add_paragraph("    ✅ 实现完整的用户认证体系")
    report.add_paragraph("    ✅ 支持论坛帖子的CRUD操作")
    report.add_paragraph("    ✅ 提供实时聊天功能")
    report.add_paragraph("    ✅ 集成资讯浏览模块")
    report.add_paragraph("    ✅ 支持Linux服务器部署")
    report.add_paragraph("    ✅ 局域网内多设备访问")
    
    # ===== 二、技术路线 =====
    report.add_section()
    report.add_title("二、技术路线", level=1)
    
    report.add_paragraph("2.1 技术架构")
    report.add_paragraph("    ┌─────────────────────────────────────────┐")
    report.add_paragraph("    │           客户端 (Avalonia UI)          │")
    report.add_paragraph("    ├─────────────────────────────────────────┤")
    report.add_paragraph("    │           后端服务 (ASP.NET Core)       │")
    report.add_paragraph("    ├─────────────────────────────────────────┤")
    report.add_paragraph("    │           数据库 (MySQL)               │")
    report.add_paragraph("    └─────────────────────────────────────────┘")
    
    report.add_paragraph("2.2 核心技术栈")
    tech_table = [
        ["分类", "技术", "版本", "说明"],
        ["前端框架", "Avalonia UI", "11.x", "跨平台桌面应用框架"],
        ["后端框架", "ASP.NET Core", "9.0", "高性能Web API框架"],
        ["ORM框架", "SqlSugar", "5.x", "MySQL数据库操作"],
        ["认证方式", "JWT", "-", "JSON Web Token认证"],
        ["实时通信", "WebSocket", "-", "即时聊天功能"],
        ["UI组件", "Ursa.Controls", "2.x", "现代化UI控件库"],
        ["WebView", "Avalonia.WebView", "12.x", "内嵌网页浏览"]
    ]
    report.add_table(tech_table)
    
    report.add_paragraph("2.3 关键组件说明")
    
    report.add_paragraph("(1) JWT认证")
    report.add_paragraph("    ```csharp")
    report.add_paragraph("    // 核心认证逻辑")
    report.add_paragraph("    public async Task<LoginResponse> LoginAsync(LoginRequest request) {")
    report.add_paragraph("        var user = await _db.Queryable<User>()")
    report.add_paragraph("            .FirstAsync(u => u.UserName == request.UserName);")
    report.add_paragraph("        if (user == null || user.Password != request.Password) {")
    report.add_paragraph("            return new LoginResponse { Success = false, Message = \"用户名或密码错误\" };")
    report.add_paragraph("        }")
    report.add_paragraph("        var token = GenerateJwtToken(user);")
    report.add_paragraph("        return new LoginResponse { Success = true, Token = token };")
    report.add_paragraph("    }")
    report.add_paragraph("    ```")
    
    report.add_paragraph("(2) 论坛帖子服务")
    report.add_paragraph("    ```csharp")
    report.add_paragraph("    // 创建帖子")
    report.add_paragraph("    public async Task<PostDto> CreatePostAsync(CreatePostRequest request, int userId) {")
    report.add_paragraph("        var post = new Post {")
    report.add_paragraph("            UserId = userId,")
    report.add_paragraph("            Title = request.Title,")
    report.add_paragraph("            Content = request.Content,")
    report.add_paragraph("            CreatedTime = DateTime.UtcNow")
    report.add_paragraph("        };")
    report.add_paragraph("        await _db.Insertable(post).ExecuteCommandAsync();")
    report.add_paragraph("        return MapToDto(post);")
    report.add_paragraph("    }")
    report.add_paragraph("    ```")
    
    # ===== 三、系统实现 =====
    report.add_section()
    report.add_title("三、系统实现", level=1)
    
    report.add_paragraph("3.1 主操作界面")
    
    report.add_paragraph("(1) 登录界面")
    report.add_paragraph("    ┌─────────────────────────────┐")
    report.add_paragraph("    │         系统登录            │")
    report.add_paragraph("    ├─────────────────────────────┤")
    report.add_paragraph("    │  用户名: [___________]      │")
    report.add_paragraph("    │  密码:   [___________]      │")
    report.add_paragraph("    │                             │")
    report.add_paragraph("    │     [ 登录 ]  [ 注册 ]      │")
    report.add_paragraph("    └─────────────────────────────┘")
    
    report.add_paragraph("(2) 主界面布局")
    report.add_paragraph("    ┌─────────────────────────────────────────┐")
    report.add_paragraph("    │  [导航]  [社区] [聊天] [资讯] [我的]    │")
    report.add_paragraph("    ├─────────┬──────────────────────────────┤")
    report.add_paragraph("    │         │                              │")
    report.add_paragraph("    │  侧边栏 │      主内容区域              │")
    report.add_paragraph("    │         │                              │")
    report.add_paragraph("    └─────────┴──────────────────────────────┘")
    
    report.add_paragraph("3.2 数据库设计")
    db_table = [
        ["表名", "字段", "类型", "说明"],
        ["users", "id", "INT", "用户ID(主键)"],
        ["users", "user_name", "VARCHAR(50)", "用户名"],
        ["users", "password", "VARCHAR(100)", "密码(加密)"],
        ["posts", "id", "INT", "帖子ID"],
        ["posts", "user_id", "INT", "作者ID(外键)"],
        ["posts", "title", "VARCHAR(255)", "标题"],
        ["posts", "content", "TEXT", "正文内容"]
    ]
    report.add_table(db_table)
    
    report.add_paragraph("3.3 API接口设计")
    api_table = [
        ["模块", "接口", "方法", "说明"],
        ["认证", "/api/auth/login", "POST", "用户登录"],
        ["认证", "/api/auth/register", "POST", "用户注册"],
        ["帖子", "/api/posts", "GET", "获取帖子列表"],
        ["帖子", "/api/posts", "POST", "创建帖子"],
        ["聊天", "/api/channels", "GET", "获取会话列表"]
    ]
    report.add_table(api_table)
    
    # ===== 四、系统部署 =====
    report.add_section()
    report.add_title("四、系统部署", level=1)
    
    report.add_paragraph("4.1 本地开发环境")
    report.add_paragraph("    ```powershell")
    report.add_paragraph("    # 启动后端服务")
    report.add_paragraph("    cd Chat.Server")
    report.add_paragraph("    dotnet run --urls \"http://0.0.0.0:5002\"")
    report.add_paragraph("    ```")
    
    report.add_paragraph("4.2 Linux服务器部署")
    report.add_paragraph("    ```bash")
    report.add_paragraph("    # 解压部署包")
    report.add_paragraph("    tar -xzf chat-server-deploy.tar.gz")
    report.add_paragraph("    # 安装依赖")
    report.add_paragraph("    sudo ./install-debian.sh")
    report.add_paragraph("    ```")
    
    # ===== 五、功能测试 =====
    report.add_section()
    report.add_title("五、功能测试", level=1)
    
    test_table = [
        ["测试项", "预期结果"],
        ["用户注册", "成功创建新用户"],
        ["用户登录", "成功获取JWT令牌"],
        ["发布帖子", "帖子保存到数据库"],
        ["浏览帖子", "显示帖子列表"],
        ["点赞帖子", "点赞数增加"],
        ["发送消息", "消息实时推送"]
    ]
    report.add_table(test_table)
    
    # ===== 六、总结 =====
    report.add_section()
    report.add_title("六、总结", level=1)
    
    report.add_paragraph("    本项目成功实现了一个功能完整的社区论坛与即时通讯系统，涵盖用户认证、论坛管理、即时聊天、资讯浏览等核心功能。系统采用前后端分离架构，支持跨平台部署，具备良好的扩展性和安全性。")
    
    # 保存文档
    output_path = "实验报告_完整版.docx"
    report.save(output_path)

def read_docx(file_path):
    """读取Word文档内容"""
    if not os.path.exists(file_path):
        print(f"文件不存在: {file_path}")
        return
    
    doc = Document(file_path)
    print("=" * 60)
    print(f"文档内容: {file_path}")
    print("=" * 60)
    
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            print(f"{i+1}. {paragraph.text}")
    
    print("=" * 60)
    
def main():
    """主函数"""
    print("=" * 60)
    print("    实验报告处理脚本")
    print("=" * 60)
    print("1. 生成新的实验报告")
    print("2. 读取现有文档")
    print("3. 退出")
    print("=" * 60)
    
    choice = input("请输入选择 (1/2/3): ")
    
    if choice == "1":
        generate_report()
    elif choice == "2":
        file_path = input("请输入文档路径: ")
        read_docx(file_path)
    elif choice == "3":
        print("退出程序")
    else:
        print("无效选择")

if __name__ == "__main__":
    main()